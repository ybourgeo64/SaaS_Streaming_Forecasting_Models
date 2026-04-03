from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_header_row(table, texts, color="1B3A5C"):
    row = table.rows[0]
    for i, txt in enumerate(texts):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, color)

def add_row(table, texts, shade=False):
    row = table.add_row()
    for i, txt in enumerate(texts):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(txt))
        run.font.size = Pt(9)
        if shade:
            set_cell_shading(cell, "EBF5FB")

def add_verdict_row(table, texts, color="D5F5E3"):
    row = table.add_row()
    for i, txt in enumerate(texts):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(txt))
        run.font.size = Pt(9)
        run.bold = True
        set_cell_shading(cell, color)

# ═══════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════
title = doc.add_heading('Kaggle Datasets vs. Recommended Public Data Sources', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Comparative Completeness Analysis for the AI-Powered SaaS & Streaming Revenue Driver Tree Model')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
run = sub.add_run('\nPrepared: March 31, 2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x85, 0x92, 0x9E)

# ═══════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=2)
doc.add_paragraph(
    'This document provides a head-to-head comparison between the two Kaggle datasets currently powering '
    'the AI-Powered SaaS & Streaming Revenue Driver Tree model and the 13 free public data sources '
    'identified in the Defensibility Improvement Plan. The comparison evaluates each source across five '
    'dimensions: Completeness, Coverage, Quality, Granularity, and Relevance.')

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Key Findings:')
run.bold = True

findings = [
    '9 of 13 recommended sources are MORE COMPLETE than the Kaggle datasets in at least one critical dimension.',
    '4 sources (Antenna, public SEC filings, OpenView/High Alpha, and ChartMogul) are substantially more complete and could serve as primary calibration anchors.',
    'No single new source can fully REPLACE both Kaggle datasets, but the combination of Antenna + SEC filings + ChartMogul covers all gaps.',
    'The Kaggle datasets retain unique value as structured, row-level training data for ML models — the new sources provide benchmark calibration, not row-level records.',
    '2 sources (Vitally, Vena Solutions) provide only marginal improvements and should be deprioritized.',
]
for f in findings:
    doc.add_paragraph(f, style='List Bullet')

# ═══════════════════════════════════════════
# 2. CURRENT KAGGLE DATASETS
# ═══════════════════════════════════════════
doc.add_heading('2. Current Kaggle Datasets in the Model', level=2)
doc.add_paragraph(
    'The revenue driver model is built on exactly two Kaggle datasets, unified at the framework level '
    '(not row-level mixed as of v2) to serve different parts of the driver tree:')

doc.add_heading('2.1 Kaggle Streaming Dataset (Acquisition & Revenue Engine)', level=3)
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
add_header_row(t, ['Attribute', 'Details'])
rows = [
    ('Role in Model', 'Acquisition cadence, pricing mix, engagement-based expansion scenarios, observed streaming MRR'),
    ('Key Observed Metrics', 'Streaming MRR (12,026), Streaming ARR (144,312), monthly new subscriptions, contract mix'),
    ('Fields Available', 'Subscription start dates, monthly_revenue (price-card proxy), plan type, household profiles, devices, usage frequency'),
    ('Churn Observable?', 'NO — contains only active memberships; churn is not directly observed'),
    ('Time Period', 'Jan 2023 – Nov 2024 (approx. 23 months of observed data)'),
    ('Sample Size', 'Cross-sectional active member snapshot (exact N not disclosed; acquisition model shows ~20-50 new subs/month)'),
    ('Geographic Coverage', 'Not specified — appears to be single-market or synthetic'),
    ('Known Limitations', 'Rate-card proxy (not billing data), no churn labels, no marketing/CAC data, no real enterprise seats'),
]
for i, (a, d) in enumerate(rows):
    add_row(t, [a, d], shade=(i % 2 == 0))

doc.add_heading('2.2 Kaggle Customer Churn Dataset (Retention Engine)', level=3)
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
add_header_row(t, ['Attribute', 'Details'])
rows = [
    ('Role in Model', 'Calibrated retention risk engine, renewal-risk views, value-at-risk segmentation, cohort hazard curves'),
    ('Key Observed Metrics', 'Churn labels (binary), contract length, tenure, support calls, payment delay, usage frequency'),
    ('Fields Available', 'Age, Gender, Tenure, Usage Frequency, Support Calls, Payment Delay, Contract Length (Monthly/Quarterly/Annual), Subscription Type (Basic/Standard/Premium)'),
    ('Churn Observable?', 'YES — explicit churn labels per customer'),
    ('Churn Rates Observed', 'Monthly: 51.6%, Quarterly: 44.0%, Annual: 46.2% (dramatically higher than real-world streaming)'),
    ('Retention AUC', 'Weighted AUC 0.716 (v2); strict AUC 0.583 (v1)'),
    ('Sample Size', 'Cross-sectional snapshot (not time-series); exact N not disclosed'),
    ('Known Limitations', 'Cross-sectional (no longitudinal tracking), no reactivation events, no expansion/upgrade events, no marketing data, churn rates unrealistically high vs. real streaming (2-7%)'),
]
for i, (a, d) in enumerate(rows):
    add_row(t, [a, d], shade=(i % 2 == 0))

doc.add_heading('2.3 Combined Kaggle Data Gap Summary', level=3)
doc.add_paragraph('The two Kaggle datasets together leave the following critical gaps:')
gaps = [
    'No reactivation event data → reactivation rates are entirely fabricated/scenario-only',
    'No expansion/upgrade event history → expansion layer is scenario-only',
    'No marketing spend or CAC data → CAC is scenario-only',
    'Rate-card pricing only → MRR is a proxy, not billing-system derived',
    'Domain mismatch → churn dataset has 44-52% churn vs. real streaming at 2-7%',
    'No regional/geographic segmentation',
    'No ARPU/ARM breakdowns by segment or tier',
]
for g in gaps:
    doc.add_paragraph(g, style='List Bullet')

# ═══════════════════════════════════════════
# 3. SOURCE-BY-SOURCE COMPARISON
# ═══════════════════════════════════════════
doc.add_heading('3. Detailed Source-by-Source Comparison', level=2)
doc.add_paragraph(
    'Each of the 13 recommended sources is evaluated against the current Kaggle datasets across five dimensions. '
    'Ratings use: ★★★ = Substantially more complete, ★★ = Moderately more complete, ★ = Marginally better or comparable, '
    '○ = Less complete or not applicable.')

# Source comparison data
sources = [
    {
        'name': '1. Netflix / Disney+ / WBD SEC Filings',
        'category': 'MRR Proxy & Domain Transfer',
        'comparison': [
            ('Completeness', '★★★', 'Provides ARPU/ARM, subscriber counts by region, revenue breakdowns, churn indicators, tier migration data. Kaggle has none of these real-world financial metrics.'),
            ('Coverage', '★★★', 'Netflix: 280M+ global subs across 4 regions. Disney+: 150M+ subs. vs. Kaggle\'s single-market snapshot of unknown size.'),
            ('Quality', '★★★', 'SEC-audited financial data — the gold standard. Kaggle data is unaudited, synthetic-looking, and of unknown provenance.'),
            ('Granularity', '★★', 'Regional and segment-level breakdowns (UCAN, EMEA, LATAM, APAC; ad-supported vs. premium). Kaggle has no geographic or tier detail.'),
            ('Relevance', '★★★', 'Directly from streaming companies — maximally relevant to streaming revenue modeling.'),
        ],
        'verdict': 'SUBSTANTIALLY MORE COMPLETE. Essential calibration anchor. Cannot replace Kaggle as ML training data but provides the ground-truth benchmarks the Kaggle data lacks.',
        'verdict_color': 'D5F5E3',
        'action': 'USE AS PRIMARY CALIBRATION — benchmark all streaming metrics against SEC filings',
    },
    {
        'name': '2. Antenna Streaming Analytics (Public Insights)',
        'category': 'Reactivation & Domain Transfer',
        'comparison': [
            ('Completeness', '★★★', 'Provides gross/net churn, resubscription curves (10%→23%→37%→41%), serial churner data, platform-specific metrics. Kaggle has ZERO reactivation or resubscription data.'),
            ('Coverage', '★★★', 'Covers all major SVOD platforms (Netflix, Disney+, Max, Peacock, etc.) with weighted averages. Kaggle covers one unidentified dataset.'),
            ('Quality', '★★★', 'Industry-recognized analytics provider used by media companies for planning. Kaggle data source is unknown.'),
            ('Granularity', '★★★', 'Platform-specific, monthly cadence, gross vs. net churn decomposition, resubscription time-decay curves. Kaggle has only binary churn labels.'),
            ('Relevance', '★★★', 'THE most directly relevant source for streaming churn and reactivation modeling.'),
        ],
        'verdict': 'SUBSTANTIALLY MORE COMPLETE — HIGHEST IMPACT SOURCE. The resubscription curve alone fills the model\'s single biggest gap (fabricated reactivation rates).',
        'verdict_color': 'ABEBC6',
        'action': 'CRITICAL ADDITION — replace fabricated reactivation with Antenna time-decay curve',
    },
    {
        'name': '3. Recurly Research (Churn & Revenue Benchmarks)',
        'category': 'MRR Proxy & Reactivation',
        'comparison': [
            ('Completeness', '★★', 'Adds involuntary churn (0.86%), churn by ARPC band, billing failure rates (15% decline risk), pause/win-back data (75% return). Kaggle has no involuntary churn distinction or billing data.'),
            ('Coverage', '★★', 'Based on 2,000+ subscription businesses across Recurly\'s platform. Broader than single Kaggle dataset but not streaming-specific.'),
            ('Quality', '★★', 'Based on real billing system data from Recurly\'s payment platform. More authoritative than Kaggle for subscription economics.'),
            ('Granularity', '★★', 'Churn by ARPC band (<$10 to >$250), voluntary vs. involuntary split, industry-specific rates. Kaggle has no revenue-band segmentation.'),
            ('Relevance', '★★', 'SaaS/subscription-focused rather than pure streaming. Good for billing failure discount and pause-based reactivation modeling.'),
        ],
        'verdict': 'MODERATELY MORE COMPLETE. Fills specific gaps (billing failure discount, involuntary churn, pause reactivation) that Kaggle cannot address.',
        'verdict_color': 'D5F5E3',
        'action': 'SUPPLEMENT — use for billing failure discount factor and involuntary churn calibration',
    },
    {
        'name': '4. ChartMogul SaaS Benchmarks Report',
        'category': 'MRR, Reactivation & Expansion',
        'comparison': [
            ('Completeness', '★★★', 'MRR growth benchmarks by ARR band, net/gross MRR churn rates, expansion MRR as % of total, reactivation MRR = ~10% of ARR gained. Kaggle has none of these composition metrics.'),
            ('Coverage', '★★', 'Based on 2,100+ SaaS companies connected to ChartMogul. Broader SaaS coverage than Kaggle but not streaming-specific.'),
            ('Quality', '★★★', 'Derived from actual billing system integrations (Stripe, Braintree, etc.) — real revenue data, not surveys.'),
            ('Granularity', '★★', 'Segmented by ARR band (<$300K to >$15M), net vs. gross churn, expansion vs. contraction MRR. Kaggle has no ARR-band segmentation.'),
            ('Relevance', '★★', 'SaaS-focused. Essential for the SaaS side of the driver tree; less relevant for streaming-specific metrics.'),
        ],
        'verdict': 'SUBSTANTIALLY MORE COMPLETE for SaaS metrics. The "reactivation = 10% of ARR gained" and expansion MRR benchmarks directly fill two fabricated layers.',
        'verdict_color': 'D5F5E3',
        'action': 'USE AS PRIMARY SaaS CALIBRATION — anchor MRR composition, expansion, and reactivation',
    },
    {
        'name': '5. KeyBanc (KBCM) Annual SaaS Survey',
        'category': 'CLV/CAC',
        'comparison': [
            ('Completeness', '★★★', 'Provides CAC ratios ($1.20 blended, $1.78 new), payback periods, S&M as % of revenue (33-47%), gross margins (77-81%), gross dollar churn. Kaggle has ZERO cost/spend data.'),
            ('Coverage', '★★', 'Based on hundreds of private SaaS companies. Enterprise and SMB segments represented.'),
            ('Quality', '★★★', 'Industry gold-standard SaaS survey, cited by VCs, CFOs, and analysts. Far more authoritative than Kaggle for unit economics.'),
            ('Granularity', '★★', 'Segmented by ACV band, backing type (PE vs. VC vs. bootstrapped), company stage. Kaggle has no segmentation on economics.'),
            ('Relevance', '★★', 'Purely SaaS-focused. Not directly applicable to streaming but essential for SaaS CAC/CLV calibration.'),
        ],
        'verdict': 'SUBSTANTIALLY MORE COMPLETE for unit economics. The only way to move CAC from "scenario-only" to "benchmark-bounded."',
        'verdict_color': 'D5F5E3',
        'action': 'SUPPLEMENT — anchor CAC sensitivity ranges to KBCM benchmarks',
    },
    {
        'name': '6. OpenView / High Alpha 2024 Benchmarks',
        'category': 'CLV/CAC & Expansion',
        'comparison': [
            ('Completeness', '★★★', 'CAC ratio ($2.00 new), expansion ARR = 40% of new ARR, NRR at 110% for public SaaS, Rule of 40. Kaggle has no expansion event data or cost metrics.'),
            ('Coverage', '★★★', '800+ companies surveyed, representing the broadest SaaS benchmark in the market.'),
            ('Quality', '★★★', 'Annual survey with 5+ years of trend data. Widely cited benchmark report.'),
            ('Granularity', '★★', 'By company size (>$50M ARR: expansion = 50-67% of new ARR), stage, and growth rate. Kaggle has none.'),
            ('Relevance', '★★', 'SaaS-focused. The "expansion = 40% of new ARR" finding is directly actionable for the expansion scenario layer.'),
        ],
        'verdict': 'SUBSTANTIALLY MORE COMPLETE for expansion and CAC. This is the single best source for converting the expansion layer from fabricated to benchmark-calibrated.',
        'verdict_color': 'D5F5E3',
        'action': 'CRITICAL ADDITION — replace fabricated expansion rates with NRR decomposition benchmarks',
    },
    {
        'name': '7. Benchmarkit 2024 B2B SaaS Metrics Report',
        'category': 'CLV/CAC',
        'comparison': [
            ('Completeness', '★★', 'CLTV:CAC ratio of 3.6, blended CAC trends, NRR median of 101%, growth rate benchmarks. Overlaps with OpenView/KBCM but adds CLTV:CAC ratio.'),
            ('Coverage', '★★', 'Private B2B SaaS companies. Good mid-market representation.'),
            ('Quality', '★★', 'Credible but less widely cited than KBCM or OpenView. Good validation cross-reference.'),
            ('Granularity', '★', 'Less granular than KBCM or OpenView. Primarily aggregate medians.'),
            ('Relevance', '★★', 'Useful specifically for the CLTV:CAC sanity check (3.6 median, 2.5-5.0 normal range).'),
        ],
        'verdict': 'MODERATELY MORE COMPLETE. Adds the CLTV:CAC ratio benchmark (3.6) which the Kaggle data cannot provide, but overlaps heavily with Sources 5 and 6.',
        'verdict_color': 'D5F5E3',
        'action': 'SUPPLEMENT (LOWER PRIORITY) — use as cross-validation for KBCM and OpenView data',
    },
    {
        'name': '8. Netflix Public Win-Back Statements',
        'category': 'Reactivation',
        'comparison': [
            ('Completeness', '★★', 'Netflix-specific: 50% return within 6 months, 61% within 1 year, 50M net gain from password crackdown. Kaggle has zero reactivation data.'),
            ('Coverage', '★', 'Single company (Netflix). Best-in-class, not representative of industry average.'),
            ('Quality', '★★★', 'Direct from Netflix executive statements in SEC-filed earnings transcripts. Unimpeachable source authority.'),
            ('Granularity', '★', 'High-level aggregate reactivation rates only. No cohort or segment breakdowns.'),
            ('Relevance', '★★', 'Directly relevant to streaming reactivation but represents upper-bound (Netflix is the retention leader).'),
        ],
        'verdict': 'MODERATELY MORE COMPLETE for reactivation. Provides upper-bound calibration point. Best used alongside Antenna\'s broader platform coverage.',
        'verdict_color': 'D5F5E3',
        'action': 'SUPPLEMENT — use as upper-bound reactivation calibration alongside Antenna',
    },
    {
        'name': '9. Churnkey Streaming Churn Benchmark Report',
        'category': 'Domain Transfer',
        'comparison': [
            ('Completeness', '★★', 'Platform-specific churn (Netflix 2.1%, Amazon 3.5-4.1%, Max 6-6.9%, Disney+ 5.5%, Apple TV+ 6.5-7%). Kaggle churn data is 44-52% — unrealistically high.'),
            ('Coverage', '★★', 'Covers 6+ streaming platforms individually. Broader platform coverage than Kaggle\'s single-source churn.'),
            ('Quality', '★★', 'Aggregated from multiple sources. Blog format but well-cited. Less authoritative than Antenna or SEC filings.'),
            ('Granularity', '★★', 'Platform-specific, bundling impact (29% vs. 70% churn potential), serial churner segmentation (23% of audience), content-completion churn (26%).'),
            ('Relevance', '★★★', 'Directly addresses the domain transfer gap — provides the real-world streaming churn rates needed to calibrate the Kaggle churn model.'),
        ],
        'verdict': 'MODERATELY MORE COMPLETE. Valuable for domain transfer calibration but somewhat redundant with Antenna + SEC filings.',
        'verdict_color': 'D5F5E3',
        'action': 'SUPPLEMENT — use for platform-specific churn mapping table alongside Antenna',
    },
    {
        'name': '10. Public SaaS Company NRR Disclosures (SEC)',
        'category': 'Expansion',
        'comparison': [
            ('Completeness', '★★', 'Company-specific NRR: Snowflake 158%, Datadog 130%, Slack >140%, HubSpot 103%, Zoom 98%. Kaggle has no NRR or expansion data.'),
            ('Coverage', '★★', 'Covers ~20+ public SaaS companies across segments (infrastructure, vertical, SMB, enterprise).'),
            ('Quality', '★★★', 'SEC-audited financial disclosures. Gold standard for expansion benchmarks.'),
            ('Granularity', '★', 'Company-level NRR only. No cohort or product-line breakdowns publicly available.'),
            ('Relevance', '★★', 'SaaS-focused. Useful for defining expansion scenario ranges (SMB vs. mid-market vs. enterprise).'),
        ],
        'verdict': 'MODERATELY MORE COMPLETE for expansion. Provides real NRR decomposition data that Kaggle completely lacks, but overlaps with OpenView aggregated data.',
        'verdict_color': 'D5F5E3',
        'action': 'SUPPLEMENT — use to define expansion scenario tiers (SMB/mid-market/enterprise)',
    },
    {
        'name': '11. Streaming Plan Upgrade/Pricing Data',
        'category': 'Expansion (Streaming)',
        'comparison': [
            ('Completeness', '★★', 'Tier pricing structures, ad-supported adoption (55% of Netflix new signups), ARPU growth (Netflix $10.82→$11.70 = 8.1% YoY). Kaggle has rate-card pricing only.'),
            ('Coverage', '★', 'Major streamers only (Netflix, Disney+, Max). Not comprehensive across all platforms.'),
            ('Quality', '★★', 'Mix of public filings (high quality) and pricing page observations (good quality).'),
            ('Granularity', '★', 'Aggregate ARPU growth and tier adoption only. No subscriber-level migration data.'),
            ('Relevance', '★★★', 'Directly addresses streaming expansion — the ad-tier→premium upgrade path is the main streaming expansion motion.'),
        ],
        'verdict': 'MODERATELY MORE COMPLETE for streaming expansion specifically. Fills a gap Kaggle cannot address (tier migration dynamics).',
        'verdict_color': 'D5F5E3',
        'action': 'SUPPLEMENT — use to model streaming-specific expansion (tier upgrade rates)',
    },
    {
        'name': '12. Vitally SaaS Churn Benchmarks',
        'category': 'Retention Calibration',
        'comparison': [
            ('Completeness', '★', 'B2B SaaS churn: 3.5% monthly avg; voluntary 2.6%, involuntary 0.8%. Churn by ARPU band. Similar data already available from Recurly.'),
            ('Coverage', '★', 'General SaaS market. No streaming coverage.'),
            ('Quality', '★', 'Blog-based aggregation. Less authoritative than Recurly or ChartMogul.'),
            ('Granularity', '★', 'ARPU-band segmentation (>$250: 5.0% vs. $25-$50: 7.3%). Useful but not unique.'),
            ('Relevance', '★', 'SaaS-only. Largely redundant with Recurly + ChartMogul data.'),
        ],
        'verdict': 'MARGINALLY MORE COMPLETE. Provides overlapping data already available from higher-quality sources (Recurly, ChartMogul). DEPRIORITIZE.',
        'verdict_color': 'FADBD8',
        'action': 'DEPRIORITIZE — data is largely redundant with Sources 3 and 4',
    },
    {
        'name': '13. Vena Solutions SaaS Churn Analysis',
        'category': 'Retention Calibration',
        'comparison': [
            ('Completeness', '★', 'SaaS churn rate analysis by segment. Similar content to Vitally and Recurly.'),
            ('Coverage', '★', 'General B2B SaaS. No streaming coverage.'),
            ('Quality', '★', 'Blog-based analysis. Lower authority than primary benchmark providers.'),
            ('Granularity', '★', 'SMB 3-7% monthly; Enterprise ≤1% monthly. Segment-level only.'),
            ('Relevance', '★', 'SaaS-only. Redundant with Recurly + ChartMogul + Vitally.'),
        ],
        'verdict': 'MARGINALLY MORE COMPLETE. Lowest incremental value of all 13 sources. DEPRIORITIZE.',
        'verdict_color': 'FADBD8',
        'action': 'DEPRIORITIZE — data is available from more authoritative sources',
    },
]

# Render each source comparison
for src in sources:
    doc.add_heading(f'Source {src["name"]}', level=3)
    p = doc.add_paragraph()
    run = p.add_run(f'Model Gap Addressed: ')
    run.bold = True
    run = p.add_run(src['category'])

    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    add_header_row(t, ['Dimension', 'Rating', 'Comparison vs. Kaggle Data'])
    for i, (dim, rating, detail) in enumerate(src['comparison']):
        add_row(t, [dim, rating, detail], shade=(i % 2 == 0))

    # Verdict
    add_verdict_row(t, ['VERDICT', '', src['verdict']], color=src['verdict_color'])
    add_verdict_row(t, ['ACTION', '', src['action']], color='D6EAF8')

    doc.add_paragraph('')  # spacer

# ═══════════════════════════════════════════
# 4. AGGREGATE COMPARISON MATRIX
# ═══════════════════════════════════════════
doc.add_heading('4. Aggregate Completeness Comparison Matrix', level=2)
doc.add_paragraph(
    'The following matrix summarizes how each source compares to the Kaggle datasets. '
    'Sources are ranked by overall completeness advantage.')

t = doc.add_table(rows=1, cols=8)
t.style = 'Table Grid'
add_header_row(t, ['#', 'Source', 'Complete-\nness', 'Cover-\nage', 'Quality', 'Granu-\nlarity', 'Rele-\nvance', 'Overall\nVerdict'])

matrix_data = [
    ('1', 'Antenna Insights', '★★★', '★★★', '★★★', '★★★', '★★★', 'CRITICAL'),
    ('2', 'SEC Filings (Netflix/Disney/WBD)', '★★★', '★★★', '★★★', '★★', '★★★', 'CRITICAL'),
    ('3', 'OpenView/High Alpha 2024', '★★★', '★★★', '★★★', '★★', '★★', 'CRITICAL'),
    ('4', 'ChartMogul Reports', '★★★', '★★', '★★★', '★★', '★★', 'HIGH'),
    ('5', 'KeyBanc (KBCM) Survey', '★★★', '★★', '★★★', '★★', '★★', 'HIGH'),
    ('6', 'Recurly Research', '★★', '★★', '★★', '★★', '★★', 'HIGH'),
    ('7', 'Churnkey Blog', '★★', '★★', '★★', '★★', '★★★', 'MEDIUM'),
    ('8', 'Netflix Win-Back Statements', '★★', '★', '★★★', '★', '★★', 'MEDIUM'),
    ('9', 'Public NRR Disclosures', '★★', '★★', '★★★', '★', '★★', 'MEDIUM'),
    ('10', 'Streaming Pricing Data', '★★', '★', '★★', '★', '★★★', 'MEDIUM'),
    ('11', 'Benchmarkit Report', '★★', '★★', '★★', '★', '★★', 'LOW'),
    ('12', 'Vitally Benchmarks', '★', '★', '★', '★', '★', 'DEPRI.'),
    ('13', 'Vena Solutions', '★', '★', '★', '★', '★', 'DEPRI.'),
]
for i, row_data in enumerate(matrix_data):
    add_row(t, row_data, shade=(i % 2 == 0))

# ═══════════════════════════════════════════
# 5. GAP-BY-GAP ANALYSIS
# ═══════════════════════════════════════════
doc.add_heading('5. Which Kaggle Gaps Each Source Fills', level=2)

t = doc.add_table(rows=1, cols=7)
t.style = 'Table Grid'
add_header_row(t, ['Kaggle Gap', 'Severity', 'Best New Source', 'What It Adds', 'Replace or\nSupplement?', 'Confidence\nImpact', 'Priority'])

gap_rows = [
    ('Fabricated reactivation rates', 'CRITICAL', 'Antenna Insights', 'Time-decay resubscription curve (10%→23%→37%→41%)', 'REPLACE the fabricated rate', '+5 to +8 pts', 'P1 (Week 1)'),
    ('Fabricated expansion rates', 'CRITICAL', 'OpenView/High Alpha', 'Expansion = 40% of new ARR; NRR decomposition by segment', 'REPLACE the fabricated rate', '+3 to +5 pts', 'P1 (Week 1)'),
    ('Domain transfer gap (44-52% vs 2-7% churn)', 'HIGH', 'Antenna + SEC Filings', 'Real streaming churn by platform (Netflix 1.8%, avg 5.3%)', 'SUPPLEMENT with calibration table', '+5 pts', 'P2 (Week 2)'),
    ('No CAC/marketing data', 'HIGH', 'KBCM + OpenView', 'CAC ratios ($1.20-$2.00 per $1 ARR), S&M benchmarks', 'SUPPLEMENT with benchmark ranges', '+2 to +3 pts', 'P2 (Week 2)'),
    ('Rate-card MRR proxy', 'MEDIUM', 'SEC Filings + Recurly', 'Real ARPU/ARM ($11.70 Netflix), billing failure rate (4.5%)', 'SUPPLEMENT with discount factor', '+1 to +2 pts', 'P3 (Week 3)'),
    ('No ARPU segmentation', 'MEDIUM', 'SEC Filings + ChartMogul', 'ARPU by region, by ARR band, by tier', 'SUPPLEMENT the price-card', '+1 pt', 'P3 (Week 3)'),
    ('No churn by segment', 'LOW', 'Recurly + Churnkey', 'Churn by ARPC band, by platform, voluntary vs involuntary', 'SUPPLEMENT retention engine', '+1 pt', 'P3 (Week 4)'),
]
for i, row_data in enumerate(gap_rows):
    add_row(t, row_data, shade=(i % 2 == 0))

# ═══════════════════════════════════════════
# 6. REPLACE VS SUPPLEMENT
# ═══════════════════════════════════════════
doc.add_heading('6. Replace vs. Supplement Recommendation', level=2)

doc.add_heading('6.1 Can Any Source Fully Replace the Kaggle Datasets?', level=3)
p = doc.add_paragraph()
run = p.add_run('Short answer: No.')
run.bold = True
doc.add_paragraph(
    'The Kaggle datasets serve as row-level ML training data — they provide individual customer records '
    'with features (tenure, support calls, payment delay, usage) and labels (churn yes/no, subscription dates). '
    'None of the 13 new sources provide row-level records suitable for training supervised ML models. '
    'They provide aggregate benchmarks, industry statistics, and calibration targets.')

doc.add_paragraph(
    'However, the new sources can REPLACE specific fabricated/scenario layers:')

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
add_header_row(t, ['Model Layer', 'Current Status', 'After New Sources'])
replacements = [
    ('Reactivation rates', 'Fabricated (scenario-only)', 'REPLACED by Antenna time-decay curve → benchmark-calibrated'),
    ('Expansion rates', 'Fabricated (scenario-only)', 'REPLACED by OpenView NRR decomposition → benchmark-bounded'),
    ('CAC estimates', 'Fabricated (scenario-only)', 'UPGRADED to benchmark-bounded range (KBCM/OpenView)'),
    ('MRR proxy', 'Rate-card (transparent proxy)', 'IMPROVED with billing failure discount (Recurly) + ARPU validation (SEC)'),
    ('Churn calibration', 'Kaggle-trained (44-52%)', 'SUPPLEMENTED with domain transfer table (Antenna/SEC → 2-7% targets)'),
    ('Retention model', 'Kaggle-trained (AUC 0.716)', 'RETAINED — Kaggle data still needed for ML training; new sources calibrate outputs'),
    ('Acquisition model', 'Kaggle-trained (MAPE 15.9%)', 'RETAINED — Kaggle data still needed for time-series training'),
]
for i, (layer, current, after) in enumerate(replacements):
    add_row(t, [layer, current, after], shade=(i % 2 == 0))

doc.add_heading('6.2 Recommended Architecture: Kaggle + Benchmark Calibration Layer', level=3)
doc.add_paragraph(
    'The optimal approach is a two-layer architecture:')
doc.add_paragraph('Layer 1 (ML Training): Continue using Kaggle datasets for supervised model training (acquisition time-series, retention classifier).', style='List Bullet')
doc.add_paragraph('Layer 2 (Benchmark Calibration): Add a calibration layer that maps ML outputs to industry-realistic ranges using the new sources as calibration targets.', style='List Bullet')
doc.add_paragraph(
    'This preserves the model\'s ML capabilities while grounding all outputs in real-world benchmarks, '
    'potentially raising the confidence score from 65/100 to 85-95/100.')

# ═══════════════════════════════════════════
# 7. SOURCES LESS COMPLETE THAN KAGGLE
# ═══════════════════════════════════════════
doc.add_heading('7. Sources Less Complete Than Kaggle (Deprioritize)', level=2)
doc.add_paragraph(
    'Two of the 13 sources provide marginal or redundant value compared to the Kaggle datasets:')

t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
add_header_row(t, ['Source', 'Why Less Valuable', 'Better Alternative Already in List', 'Recommendation'])
add_row(t, ['Vitally SaaS Churn Benchmarks', 'Blog-based aggregation with data already available from Recurly and ChartMogul at higher quality and granularity', 'Recurly (Source 3) + ChartMogul (Source 4)', 'DEPRIORITIZE — remove from Week 1-2 roadmap'])
add_row(t, ['Vena Solutions Churn Analysis', 'Blog-format analysis providing SMB/Enterprise churn segments already covered by Recurly, ChartMogul, and Vitally', 'Recurly (Source 3) + ChartMogul (Source 4)', 'DEPRIORITIZE — remove from Week 1-2 roadmap'], shade=True)

doc.add_paragraph('')
doc.add_paragraph(
    'Note: These sources are not "bad" — they simply don\'t add incremental value beyond what Sources 3-6 already provide. '
    'They could serve as tertiary cross-validation references if time permits.')

# ═══════════════════════════════════════════
# 8. FINAL RECOMMENDATIONS
# ═══════════════════════════════════════════
doc.add_heading('8. Final Prioritized Recommendations', level=2)

doc.add_heading('Tier 1: Critical — More Complete Than Kaggle, Immediate Integration (Week 1)', level=3)
tier1 = [
    'Antenna Insights → Replace fabricated reactivation rates with empirical resubscription curve (+5-8 confidence points)',
    'OpenView/High Alpha → Replace fabricated expansion rates with NRR decomposition benchmarks (+3-5 confidence points)',
    'SEC Filings (Netflix/Disney/WBD) → Validate streaming ARPU and establish domain transfer calibration targets',
]
for t1 in tier1:
    doc.add_paragraph(t1, style='List Bullet')

doc.add_heading('Tier 2: High Value — More Complete, Scheduled Integration (Week 2)', level=3)
tier2 = [
    'ChartMogul → Anchor SaaS MRR composition and reactivation MRR benchmarks',
    'KeyBanc (KBCM) → Establish CAC benchmark ranges for unit economics layer',
    'Recurly → Apply billing failure discount to MRR proxy and add involuntary churn segmentation',
]
for t2 in tier2:
    doc.add_paragraph(t2, style='List Bullet')

doc.add_heading('Tier 3: Supplementary — Moderate Improvement (Week 3-4)', level=3)
tier3 = [
    'Churnkey, Netflix Win-Back, Public NRR Disclosures, Streaming Pricing Data → Cross-validate and fill specific sub-gaps',
    'Benchmarkit → Tertiary validation for CLTV:CAC ratio',
]
for t3 in tier3:
    doc.add_paragraph(t3, style='List Bullet')

doc.add_heading('Tier 4: Deprioritize — Redundant with Higher-Quality Sources', level=3)
tier4 = [
    'Vitally SaaS Churn Benchmarks — redundant with Recurly + ChartMogul',
    'Vena Solutions Churn Analysis — redundant with Recurly + ChartMogul + Vitally',
]
for t4 in tier4:
    doc.add_paragraph(t4, style='List Bullet')

doc.add_heading('8.1 Expected Confidence Score Trajectory', level=3)
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
add_header_row(t, ['Milestone', 'Sources Integrated', 'Projected Score'])
add_row(t, ['Current (Kaggle only)', 'None', '65 / 100'])
add_row(t, ['After Tier 1', 'Antenna + OpenView + SEC Filings', '73–80 / 100'], shade=True)
add_row(t, ['After Tier 1 + 2', '+ ChartMogul + KBCM + Recurly', '80–88 / 100'])
add_row(t, ['After All Tiers', 'All 11 prioritized sources', '85–95 / 100'], shade=True)

doc.save('/home/ubuntu/kaggle_vs_new_sources_comparison.docx')
print("Document saved successfully.")
