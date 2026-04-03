from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import datetime

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Title
title = doc.add_heading('Multi-Source Data Methodology: Defensibility Guidance', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0, 51, 102)

# Subtitle
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('For SaaS/Streaming Revenue Driver Models Using Public Data Sources')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)
run.italic = True

doc.add_paragraph()  # spacer

# Date
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = date_p.add_run(f'Prepared: {datetime.date.today().strftime("%B %d, %Y")}')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

# Helper
def add_body(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p

# ── SECTION 1: DIRECT ANSWER ──
doc.add_heading('1. The Direct Answer: Yes, Multi-Source Is Standard Practice', level=1)

p = doc.add_paragraph()
run = p.add_run('Using multiple specialized datasets to support a revenue driver model is not only acceptable—it is the methodological gold standard.')
run.bold = True
run.font.size = Pt(11)

add_body(
    'In financial modeling, academic research, and industry analysis, relying on a single dataset when '
    'better-fit sources exist for specific components would be the weaker methodological choice. '
    'The practice of integrating multiple data sources is formally recognized as "data triangulation" '
    'and is a cornerstone of rigorous empirical research (Denzin, 1978; Patton, 1999). '
    'It is standard practice in every major consulting firm (McKinsey, BCG, Bain), '
    'investment bank research division, and peer-reviewed academic journal.'
)

add_body(
    'The key insight: No single public dataset comprehensively covers all dimensions of a SaaS/streaming '
    'revenue model (acquisition, retention, expansion, reactivation, MRR, CLV/CAC). Using the best '
    'available source for each component produces a more defensible model than forcing one incomplete '
    'dataset to serve all purposes.'
)

# ── SECTION 2: WHY MULTI-SOURCE IS MORE DEFENSIBLE ──
doc.add_heading('2. Why Multi-Source Is More Defensible Than Single-Source', level=1)

doc.add_heading('2.1 Advantages of Specialized Sources', level=2)

advantages = [
    ('Higher data quality per component', 'Each metric draws from the source with the deepest coverage, largest sample, and most rigorous collection methodology for that specific domain.'),
    ('Reduced single-source bias', 'Any individual dataset has collection biases, sampling limitations, or coverage gaps. Multiple sources dilute these weaknesses.'),
    ('Cross-validation capability', 'Where sources overlap, you can verify consistency—strengthening confidence in the data. Divergences surface areas requiring deeper investigation.'),
    ('Transparency and auditability', 'Readers can independently verify each claim against its cited source, rather than trusting an opaque monolithic dataset.'),
    ('Reflects real-world practice', 'No operating company relies on a single data source for forecasting. Blending CRM data, billing systems, market research, and financial statements is standard.'),
]

for title_text, desc in advantages:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('2.2 When Single-Source Would Be Preferred', level=2)
add_body('A single source is preferable only when:')
add_bullet('One dataset genuinely covers all required metrics with sufficient depth and quality')
add_bullet('The analysis requires strict temporal or methodological consistency that combining sources would violate (e.g., a controlled experiment)')
add_bullet('The datasets cannot be meaningfully harmonized due to incompatible definitions or timeframes')

add_body(
    'For a revenue driver model spanning acquisition, retention, expansion, and unit economics across '
    'SaaS and streaming verticals, no single public dataset meets these criteria. '
    'Multi-source is therefore the correct approach.'
)

# ── SECTION 3: RISKS AND MITIGATIONS ──
doc.add_heading('3. Risks of Multi-Source Integration and How to Mitigate Them', level=1)

# Table
table = doc.add_table(rows=6, cols=3)
table.style = 'Medium Shading 1 Accent 1'
table.autofit = True

headers = ['Risk', 'Description', 'Mitigation']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

risks = [
    ('Definitional inconsistency', 'Different sources may define "churn" or "MRR" differently', 'Document each source\'s definition explicitly; normalize or note differences'),
    ('Temporal mismatch', 'Data from different time periods may not be comparable', 'Align time windows; clearly state the reporting period for each source'),
    ('Cherry-picking perception', 'Critics may claim you selected sources that support your thesis', 'Establish source selection criteria before analysis; document why each source was chosen'),
    ('Aggregation errors', 'Combining data at different granularities can introduce artifacts', 'Maintain component-level analysis; do not arithmetically combine incompatible figures'),
    ('Reproducibility concerns', 'Readers must be able to trace every claim to its source', 'Provide a complete data source appendix with URLs, access dates, and extraction methods'),
]

for row_idx, (risk, desc, mit) in enumerate(risks, 1):
    table.rows[row_idx].cells[0].text = risk
    table.rows[row_idx].cells[1].text = desc
    table.rows[row_idx].cells[2].text = mit

doc.add_paragraph()  # spacer

# ── SECTION 4: BEST PRACTICES ──
doc.add_heading('4. Best Practices for a Defensible Multi-Source Approach', level=1)

doc.add_heading('4.1 Source Selection Protocol', level=2)
add_body('For each model component, document the following in a "Data Source Selection Matrix":')
add_bullet('Metric needed (e.g., "Monthly churn rate for SaaS companies")')
add_bullet('Source selected and why (coverage, sample size, recency, methodology)')
add_bullet('Alternative sources considered and why they were not chosen')
add_bullet('Known limitations of the selected source')
add_bullet('How the data was accessed, extracted, and any transformations applied')

doc.add_heading('4.2 Documentation Requirements', level=2)
add_body('Your article or model appendix should include:')
add_bullet('A Data Sources Table listing every source, what it provides, its methodology, and time period')
add_bullet('A Methodology Note explaining the multi-source approach and why it was chosen')
add_bullet('Inline citations linking every quantitative claim to its specific source')
add_bullet('A Limitations section acknowledging where data is approximate, interpolated, or proxy-based')

doc.add_heading('4.3 Handling Proxies and Fabricated Data Replacement', level=2)
add_body(
    'When replacing fabricated Kaggle metrics with real public data, be explicit about the transition:'
)
add_bullet('Label proxy metrics clearly (e.g., "Industry median used as proxy for company-specific retention")')
add_bullet('State the direction and magnitude of potential bias the proxy introduces')
add_bullet('Where possible, provide sensitivity analysis showing how results change if the proxy is off by ±10-20%')
add_bullet('Never present proxy data as if it were direct measurement')

doc.add_heading('4.4 The "Minimum Defensibility" Standard', level=2)
add_body('For a LinkedIn article defending a revenue driver model, your multi-source approach is defensible if:')
add_bullet('Every quantitative claim has a cited, verifiable public source', level=0)
add_bullet('You acknowledge where real data replaces previously fabricated data', level=0)
add_bullet('You explain your source selection rationale (even one sentence per source suffices)', level=0)
add_bullet('You note known limitations honestly rather than overstating precision', level=0)
add_bullet('A knowledgeable reader could reproduce your data gathering process', level=0)

# ── SECTION 5: SPECIFIC RECOMMENDATIONS ──
doc.add_heading('5. Specific Recommendations for This Revenue Driver Model', level=1)

add_body('Given that the model spans Acquisition, Retention, Expansion, Reactivation, MRR, and CLV/CAC across SaaS and streaming verticals:')

recs = [
    'Use specialized sources for each component. A SaaS benchmarking report (e.g., OpenView, KeyBanc, Bessemer) is authoritative for retention and expansion metrics; a streaming industry report (e.g., Antenna, Parks Associates) is authoritative for subscriber dynamics. Neither alone covers both.',
    'Prioritize authoritative industry sources over Kaggle. Kaggle datasets are valuable for learning and prototyping but are rarely the most defensible source for published analysis. Government statistics, industry association reports, SEC filings, and established research firms carry more weight.',
    'Structure the article around the multi-source approach as a strength. Frame it as: "We assembled the most authoritative public data for each revenue component" rather than apologizing for not having one perfect dataset.',
    'Create a Source Authority Hierarchy: (1) Government/regulatory filings, (2) Established research firms and industry associations, (3) Company disclosures and SEC filings, (4) Academic peer-reviewed studies, (5) Reputable business publications, (6) Community datasets (Kaggle). Use the highest-authority source available for each metric.',
    'Cross-validate where sources overlap. If two sources report similar churn rates or growth figures, note the convergence—it strengthens both. If they diverge, investigate and explain why.',
]

for i, rec in enumerate(recs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. ')
    run.bold = True
    p.add_run(rec)

# ── SECTION 6: INDUSTRY PRECEDENT ──
doc.add_heading('6. Industry and Academic Precedent', level=1)

add_body('Multi-source data integration is not merely acceptable—it is expected in serious analytical work:')

precedents = [
    ('Investment Banking & Equity Research', 'Every sell-side equity research report combines company filings, industry data (IDC, Gartner, Forrester), government statistics, and proprietary surveys. No analyst would defend a model built on a single Kaggle dataset.'),
    ('Management Consulting', 'McKinsey, BCG, and Bain routinely triangulate client data, market research, expert interviews, and public datasets. The triangulation itself is considered a quality indicator.'),
    ('Academic Finance & Economics', 'Published papers in top journals (Journal of Finance, AER, QJE) routinely combine CRSP, Compustat, Census data, and survey data. Reviewers expect multi-source validation.'),
    ('SaaS Industry Benchmarking', 'Leading SaaS benchmark reports (OpenView, KeyBanc KBCM Survey, Bessemer) are themselves multi-source, combining survey data, financial filings, and operational metrics from hundreds of companies.'),
    ('Data Triangulation Methodology', 'Formalized by Denzin (1978) and widely adopted across social sciences, business research, and market analysis. The explicit use of multiple independent sources to validate findings is a recognized methodological strength, not a weakness.'),
]

for title_text, desc in precedents:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}: ')
    run.bold = True
    p.add_run(desc)

# ── SECTION 7: CONCLUSION ──
doc.add_heading('7. Conclusion', level=1)

p = doc.add_paragraph()
run = p.add_run(
    'Using 9+ specialized public data sources instead of a single Kaggle dataset is not a compromise—it '
    'is an upgrade in methodological rigor. '
)
run.bold = True
p.add_run(
    'The defensibility of your revenue driver model increases when each component draws from '
    'the most authoritative available source for that specific metric, provided you document your '
    'source selection rationale, acknowledge limitations, and maintain transparent citations throughout.'
)

add_body('')
p = doc.add_paragraph()
run = p.add_run('Bottom line for the LinkedIn article: ')
run.bold = True
p.add_run(
    'Lead with the multi-source approach as a deliberate methodological choice. '
    'State that you surveyed available public data and selected the most authoritative source '
    'for each revenue component. Include a data sources table. This positions you as rigorous '
    'and transparent—exactly what a defensible analysis requires.'
)

# Save
doc.save('/home/ubuntu/multi_source_data_methodology_guidance.docx')
print("Document created successfully.")
