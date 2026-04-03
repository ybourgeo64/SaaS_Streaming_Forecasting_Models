from pathlib import Path
import math
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = Path('/home/ubuntu')
WORK = BASE / 'revenue_rebuild'
DATA = WORK / 'data'
OUT = WORK / 'output'
GUIDE = BASE / 'revenue_driver_models_user_guide.docx'
IMG = OUT / 'user_guide_assets'
IMG.mkdir(parents=True, exist_ok=True)

# Load model artifacts
param = pd.read_csv(DATA / 'parameter_registry_v4.csv')
source = pd.read_csv(DATA / 'source_registry_v4.csv')
saas_scen = pd.read_csv(DATA / 'saas_scenarios_v4.csv')
stream_scen = pd.read_csv(DATA / 'streaming_scenarios_v4.csv')
saas_base = pd.read_csv(DATA / 'saas_base_v4.csv')
stream_base = pd.read_csv(DATA / 'streaming_base_v4.csv')
mc = pd.read_csv(DATA / 'saas_clv_cac_uncertainty_v4.csv')
saas_drivers = pd.read_csv(DATA / 'saas_dominant_drivers_v4.csv')
stream_drivers = pd.read_csv(DATA / 'streaming_dominant_drivers_v4.csv')

# ----------------------------------------------------------------------
# Visuals
# ----------------------------------------------------------------------
def save_connected_system(path):
    fig, ax = plt.subplots(figsize=(12, 7))
    ax.set_xlim(0, 12); ax.set_ylim(0, 7); ax.axis('off')
    boxes = {
        'Acquisition': (0.7, 4.7, 2.0, 1.0, '#D9EAF7'),
        'Retention': (3.4, 4.7, 2.0, 1.0, '#E2F0D9'),
        'Expansion': (6.1, 4.7, 2.0, 1.0, '#FFF2CC'),
        'Reactivation': (8.8, 4.7, 2.0, 1.0, '#FCE4D6'),
        'Revenue Outputs': (4.3, 2.2, 3.2, 1.0, '#EADCF8')
    }
    for label, (x,y,w,h,color) in boxes.items():
        ax.add_patch(FancyBboxPatch((x,y), w,h, boxstyle='round,pad=0.03,rounding_size=0.08', facecolor=color, edgecolor='#1f1f1f', linewidth=1.5))
        ax.text(x+w/2, y+h/2, label, ha='center', va='center', fontsize=13, fontweight='bold')
    arrows = [((2.7,5.2),(3.4,5.2)),((5.4,5.2),(6.1,5.2)),((8.1,5.2),(8.8,5.2)),((1.7,4.7),(5.9,3.2)),((4.4,4.7),(5.9,3.2)),((7.1,4.7),(5.9,3.2)),((9.8,4.7),(5.9,3.2))]
    for a,b in arrows:
        ax.add_patch(FancyArrowPatch(a,b,arrowstyle='->',mutation_scale=15,linewidth=1.8,color='#404040'))
    labels = [
        (1.7,6.25,'New customers / subscribers'),
        (4.4,6.25,'Churn, renewals, gross retention'),
        (7.1,6.25,'Upsell, price, mix, add-ons'),
        (9.8,6.25,'Win-back and recovery'),
        (5.9,1.6,'ARR, MRR, bookings, NRR, CAC payback, net churn')
    ]
    for x,y,t in labels:
        ax.text(x,y,t,ha='center',va='center',fontsize=10)
    ax.text(6,0.6,'Connected driver forecasting treats revenue as a system, not a single top-line guess.',ha='center',fontsize=12,style='italic')
    plt.tight_layout(); fig.savefig(path,dpi=220,bbox_inches='tight'); plt.close(fig)

def save_driver_tree(path):
    fig, ax = plt.subplots(figsize=(13, 8))
    ax.set_xlim(0, 13); ax.set_ylim(0, 8); ax.axis('off')
    nodes = [
        ('Top-of-funnel',0.8,6.2),('Acquisition rate',0.8,4.8),('New logos / adds',0.8,3.4),
        ('Retention',4.0,6.2),('Renewals',4.0,4.8),('Net churn',4.0,3.4),
        ('Expansion',7.2,6.2),('Pricing / ARPU / ARPA',7.2,4.8),('Reactivation',7.2,3.4),
        ('Revenue Outcomes',10.2,5.0),('ARR / MRR / NRR / CLV',10.2,3.2)
    ]
    for text,x,y in nodes:
        ax.add_patch(FancyBboxPatch((x,y),2.0,0.8,boxstyle='round,pad=0.02,rounding_size=0.06',facecolor='#F7F7F7',edgecolor='#1f1f1f'))
        ax.text(x+1.0,y+0.4,text,ha='center',va='center',fontsize=11)
    arrow_pairs=[((2.8,6.6),(4.0,6.6)),((2.8,5.2),(4.0,5.2)),((2.8,3.8),(4.0,3.8)),((6.0,6.6),(7.2,6.6)),((6.0,5.2),(7.2,5.2)),((6.0,3.8),(7.2,3.8)),((9.2,5.4),(10.2,5.4)),((9.2,3.8),(10.2,3.6))]
    for a,b in arrow_pairs:
        ax.add_patch(FancyArrowPatch(a,b,arrowstyle='->',mutation_scale=15,linewidth=1.6,color='#404040'))
    ax.text(6.5,1.2,'Each metric is both an output from prior drivers and an input into later outcomes.',ha='center',fontsize=12,style='italic')
    plt.tight_layout(); fig.savefig(path,dpi=220,bbox_inches='tight'); plt.close(fig)

def save_workflow(path):
    fig, ax = plt.subplots(figsize=(12, 7))
    ax.set_xlim(0,12); ax.set_ylim(0,7); ax.axis('off')
    steps=[('1. Open workbook\nand review dashboard',0.7,5.3),('2. Update\nParameter Registry',3.0,5.3),('3. Select scenario\nassumptions',5.3,5.3),('4. Review monthly\nforecast outputs',7.6,5.3),('5. Interpret dominant\ndrivers and risks',9.9,5.3),('6. Compare to actuals\nand refine',5.3,2.4)]
    for txt,x,y in steps:
        ax.add_patch(FancyBboxPatch((x,y),1.8,1.0,boxstyle='round,pad=0.04',facecolor='#D9EAF7',edgecolor='#1f1f1f'))
        ax.text(x+0.9,y+0.5,txt,ha='center',va='center',fontsize=11)
    arrows=[((2.5,5.8),(3.0,5.8)),((4.8,5.8),(5.3,5.8)),((7.1,5.8),(7.6,5.8)),((9.4,5.8),(9.9,5.8)),((10.8,5.3),(6.2,3.4)),((5.3,3.4),(1.6,5.3))]
    for a,b in arrows:
        ax.add_patch(FancyArrowPatch(a,b,arrowstyle='->',mutation_scale=14,linewidth=1.6,color='#404040'))
    ax.text(6,0.9,'Use the models as a recurring management process: update, test, interpret, and learn.',ha='center',fontsize=12,style='italic')
    plt.tight_layout(); fig.savefig(path,dpi=220,bbox_inches='tight'); plt.close(fig)

def save_sheet_map(path):
    fig, ax = plt.subplots(figsize=(12, 7))
    ax.set_xlim(0, 12); ax.set_ylim(0, 7); ax.axis('off')
    sheets=['Dashboard','Assumptions','Data Sources','Manual Override Log','Reactivation Curves','Streaming Model','SaaS Model','Calculations','Monthly Forecast','Sensitivity Analysis','Dependency Map','Confidence']
    positions=[(0.8,5.7),(3.1,5.7),(5.4,5.7),(7.7,5.7),(10.0,5.7),(0.8,4.0),(3.1,4.0),(5.4,4.0),(7.7,4.0),(10.0,4.0),(3.1,2.3),(7.7,2.3)]
    for (x,y),name in zip(positions,sheets):
        ax.add_patch(FancyBboxPatch((x,y),1.8,0.8,boxstyle='round,pad=0.03',facecolor='#F7F7F7',edgecolor='#1f1f1f'))
        ax.text(x+0.9,y+0.4,name,ha='center',va='center',fontsize=10)
    ax.text(6,1.0,'Suggested workflow: Dashboard → Assumptions → Model tabs → Monthly Forecast → Sensitivity Analysis → Override Log',ha='center',fontsize=12,style='italic')
    plt.tight_layout(); fig.savefig(path,dpi=220,bbox_inches='tight'); plt.close(fig)

connected_png = IMG / 'connected_system.png'
driver_tree_png = IMG / 'driver_tree.png'
workflow_png = IMG / 'workflow.png'
sheet_map_png = IMG / 'sheet_map.png'
save_connected_system(connected_png)
save_driver_tree(driver_tree_png)
save_workflow(workflow_png)
save_sheet_map(sheet_map_png)

# ----------------------------------------------------------------------
# Word helpers
# ----------------------------------------------------------------------
def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'),'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'),'preserve'); instrText.text='PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)

def add_toc(paragraph):
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar'); fldChar.set(qn('w:fldCharType'),'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'),'preserve'); instrText.text='TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'separate')
    txt = OxmlElement('w:t'); txt.text='Right-click and update field to generate the table of contents.'
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'),'end')
    run._r.append(fldChar); run._r.append(instrText); run._r.append(fldChar2); run._r.append(txt); run._r.append(fldChar3)


def add_df_table(doc, df, title=None, max_rows=None):
    if title:
        doc.add_heading(title, level=3)
    d = df.head(max_rows) if max_rows else df
    table = doc.add_table(rows=1, cols=len(d.columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(d.columns):
        p = table.rows[0].cells[i].paragraphs[0]
        run = p.add_run(str(c)); run.bold = True
    for row in d.itertuples(index=False):
        cells = table.add_row().cells
        for i, v in enumerate(row):
            if isinstance(v, float):
                txt = f'{v:,.4f}' if abs(v) < 1 else f'{v:,.2f}'
                if float(v).is_integer():
                    txt = f'{v:,.0f}'
            else:
                txt = str(v)
            cells[i].text = txt
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

# ----------------------------------------------------------------------
# Build document
# ----------------------------------------------------------------------
doc = Document()
doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].font.size = Pt(10.5)
for style_name in ['Heading 1','Heading 2','Heading 3']:
    doc.styles[style_name].font.name = 'Arial'
footer = doc.sections[0].footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_page_number(footer)

# Cover
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Revenue Driver Models User Guide\n'); r.bold = True; r.font.size = Pt(22)
p.add_run('Practical instructions for using the SaaS and Streaming forecasting models').font.size = Pt(13)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Prepared for operating teams, finance leaders, revenue leaders, and strategy teams').italic = True
doc.add_paragraph('April 1, 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# TOC
doc.add_heading('Table of Contents', level=1)
add_toc(doc.add_paragraph())
doc.add_page_break()

# 1 Executive Summary
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'These revenue driver models are practical forecasting tools designed to show how recurring revenue is created, protected, and expanded over time. Rather than starting with a top-line revenue number and working backward, the models start with the underlying drivers that create recurring revenue: acquisition, retention, expansion, reactivation, pricing, capacity, and timing. The result is a forecast that is easier to explain, easier to challenge, and easier to update as conditions change.'
)
doc.add_paragraph(
    'The central forecasting thesis behind the models is simple: revenue behaves like a connected system. New customers or subscribers enter the system through acquisition. Some renew, some churn, some expand, and some return through reactivation. Pricing changes can increase monetization but also increase churn or reduce conversion. Capacity bottlenecks can slow growth even when demand exists. Because these effects interact, no single metric tells the full story. Connected driver forecasting matters because it preserves those cause-and-effect relationships.'
)
doc.add_paragraph(
    'The models are intended for teams that need a disciplined, transparent, and scenario-driven way to plan recurring revenue. Typical users include finance teams building annual plans, revenue operations teams managing the pipeline-to-revenue process, customer success teams looking at renewals and retention, pricing teams testing monetization changes, strategy teams evaluating downside risk, and executive teams preparing board or investor materials. Each user can start from the same workbook and still answer different business questions.'
)
doc.add_paragraph(
    'In the SaaS model, you can forecast customer growth, ARR, MRR, bookings, renewals, expansion, contraction, reactivation, CAC payback, and unit economics over a 24-month horizon. In the Streaming model, you can forecast subscribers, MRR, ARR run rate, churn, reactivation, ARPU dynamics, billing leakage, and monetization sensitivity over the same time horizon. Both models are designed to support monthly operating reviews as well as annual planning cycles.'
)
doc.add_paragraph(
    'The biggest benefit of using these models is not just a better number. It is a better explanation. When forecast outcomes change, the workbook makes it possible to trace the reason: slower acquisition, higher churn, weaker price realization, lower expansion, longer implementation delays, tighter capacity, or a tougher external environment. That clarity helps management teams focus on the few levers that matter most instead of reacting to surface-level variance.'
)
doc.add_paragraph(
    'Use cases include annual budgeting, territory or segment planning, pricing and packaging reviews, investor or board scenario preparation, churn-reduction planning, acquisition planning, capacity management, and operating cadence reviews. The models are also useful as communication tools. Teams can align around a common framework, document assumption changes in one place, and compare actual performance against the original driver logic rather than against an unexplained target.'
)
doc.add_paragraph(
    'A practical way to think about the models is this: they are management tools for asking “what has to be true?” If you want to reach a certain ARR outcome, what acquisition, retention, and expansion levels would have to hold? If you are worried about a recession or a pricing change, how much damage does that create and which drivers are affected first? If your forecast misses, which assumptions should be revised? The workbook is built to answer those questions in a disciplined way.'
)
doc.add_picture(str(workflow_png), width=Inches(6.8))
doc.add_paragraph('Figure 1. Suggested forecasting workflow for both models.')

# 2 Understanding connected driver system
doc.add_heading('2. Understanding the Connected Driver System', level=1)
doc.add_paragraph(
    'The core concept behind both models is the connected driver system: Acquisition → Retention → Expansion → Reactivation. Acquisition introduces new customers or subscribers into the system. Retention determines how much of that base stays active. Expansion measures the additional revenue created from the retained base through pricing, upsell, mix shift, or add-ons. Reactivation captures the portion of previously lost customers or subscribers who return. Together, these flows determine recurring revenue.'
)
doc.add_paragraph(
    'Traditional forecasts often isolate metrics that should not be isolated. For example, a team may set an acquisition target without adjusting for implementation capacity. Or it may plan a price increase without reflecting the likely effect on churn and conversion. Or it may assume a retention level without considering the impact of new-customer mix, customer age, or competitive pressure. Those isolated assumptions create forecasts that look neat in presentation decks but break under operating pressure. Connected driver forecasting corrects this by forcing assumptions to interact.'
)
doc.add_paragraph(
    'In a connected system, every driver has an upstream and downstream effect. A higher acquisition rate boosts the active base, but it can also change the age mix of the portfolio and therefore future churn behavior. Stronger retention supports renewal revenue, preserves expansion opportunities, and reduces the burden on acquisition. Higher expansion can lift ARR or MRR, but aggressive pricing may also create additional churn if elasticity is ignored. Reactivation offsets some churn, which changes net churn and strengthens revenue continuity. This is why the models treat forecasting as a system rather than a linear spreadsheet.'
)
doc.add_paragraph(
    'The fourteen SaaS metrics in the driver tree are not independent scorecards. ACV influences CAC. CAC affects payback and efficient growth. Renewals and retention determine the size of the revenue base that can expand. Expansion and contraction shape NRR and net churn. Seats or licenses influence contract size, ACV, and revenue scaling. Net new subscriptions depend on both acquisition and churn, not acquisition alone. That interconnection is the reason the models are useful for diagnosis as well as forecasting.'
)
doc.add_paragraph(
    'Scenario planning becomes especially powerful when the system is connected. A recession scenario can reduce acquisition, increase churn, and slow expansion simultaneously. A competitive scenario can suppress pricing power and increase customer movement. A failed pricing scenario may raise nominal ARPU but damage conversion or retention enough to reduce the long-term result. Because the models preserve these relationships, scenario analysis becomes a strategic tool rather than a cosmetic exercise.'
)
doc.add_picture(str(connected_png), width=Inches(6.8))
doc.add_paragraph('Figure 2. Core connected driver system for both models.')
doc.add_picture(str(driver_tree_png), width=Inches(6.8))
doc.add_paragraph('Figure 3. Revenue outcomes depend on interacting upstream drivers.')
doc.add_paragraph(
    'In practice, teams should think about three layers of forecasting. The first layer is the operational layer: what is changing in acquisition, churn, pricing, or capacity this quarter? The second layer is the revenue layer: how do those changes flow through ARR, MRR, bookings, or subscriber revenue? The third layer is the decision layer: which actions should management prioritize? The workbook is designed to connect all three layers so the forecast is not just a number but a management conversation.'
)

# 3 14 SaaS metrics
doc.add_heading('3. The 14 SaaS Metrics Explained', level=1)
doc.add_paragraph('This section explains the core SaaS metrics that appear in the driver forecast tree. For each metric, the goal is not only to define it, but to show where it sits in the system: what drives it, what it drives next, and how to read it intelligently in a forecast.')
metrics = [
('1. ACV (Annual Contract Value)',
 'ACV represents the annualized contract value associated with a typical customer relationship. In the model, ACV is derived primarily from ARPA multiplied by twelve, so it serves as a bridge between monthly account value and annual contract economics.',
 'ACV is driven by pricing, package mix, seat or license volume, and customer segment. It drives CAC benchmarking, payback analysis, and how new-logo additions translate into ARR. When ACV rises, the forecast may strengthen, but only if conversion and retention hold. A common pitfall is to raise ACV in the forecast without also asking whether higher pricing changes win rates or churn.'),
('2. TCV (Total Contract Value)',
 'TCV represents the total contract value over the assumed contract term. In the workbook it is used as a simple extension of ACV to provide contract-size context, especially for planning larger or longer agreements.',
 'TCV is driven by ACV and by assumed contract duration. It helps teams think about deal quality, sales prioritization, and the size of enterprise opportunities. It does not replace recurring-revenue analysis, because a large TCV number can still hide poor renewal economics. A common forecasting mistake is to let TCV dominate the planning conversation when recurring value and renewal quality matter more.'),
('3. ARR (Annual Recurring Revenue)',
 'ARR is the main annualized recurring revenue measure in the SaaS model. It is built from the customer and revenue bridges: retained ARR after churn, reactivated ARR, new-logo ARR, and expansion minus contraction.',
 'ARR is driven by acquisition, renewals, churn, reactivation, pricing, seat growth, and expansion dynamics. It drives board reporting, annual planning, valuation narratives, and operating targets. Interpret ARR as a system output rather than a goal by itself. A common pitfall is to forecast ARR directly without understanding whether the underlying acquisition and retention assumptions are feasible.'),
('4. MRR (Monthly Recurring Revenue)',
 'MRR is the monthly recurring revenue equivalent of ARR and is calculated as ARR divided by twelve. It is useful for near-term operational forecasting and monthly trend analysis.',
 'MRR is driven by the same mechanics as ARR, but it is often easier to connect to monthly operating rhythms like new bookings, churn, or price actions. It drives monthly forecasting, cash-flow awareness, and operating reviews. The pitfall is treating MRR as simply an accounting output instead of using it to see when growth accelerates, stalls, or becomes riskier.'),
('5. Subscription Bookings',
 'Subscription bookings represent the annual recurring value of new contracts booked during the period. In the model, bookings are tied to recognized new-logo activity after sales-cycle and onboarding timing are considered.',
 'Bookings are driven by pipeline generation, close rates, sales capacity, deal timing, pricing, and enterprise lumpiness. Bookings drive future ARR growth and provide an early signal of momentum. The pitfall is equating bookings with recognized recurring revenue too quickly. The model separates booking flow from recognized revenue timing for that reason.'),
('6. Renewals',
 'Renewals represent the portion of the customer base that remains active after churn. In the model, renewals are the retained customer base before reactivation and new-logo additions are layered back in.',
 'Renewals are driven by retention quality, product value, customer success, competitive intensity, and price response. They drive retained ARR, future expansion potential, and revenue stability. The pitfall is looking only at renewal counts without asking whether the retained base is healthy enough to expand or whether it is shrinking in value.'),
('7. Reactivations',
 'Reactivations capture previously churned customers who return. The model uses explicit return curves so that win-back behavior happens over time rather than instantly.',
 'Reactivations are driven by churn pools, win-back effectiveness, billing recovery, product changes, and timing. They support lower net churn and ARR continuity, and more realistic customer-base evolution. The pitfall is counting reactivations as fully new growth or assuming they happen immediately after churn. The workbook treats them as a distinct flow.'),
('8. CAC (Customer Acquisition Cost)',
 'CAC is the cost to acquire a new customer. In the workbook, CAC is estimated independently rather than backed into the answer, and it is linked to ACV and scenario conditions.',
 'CAC is driven by sales efficiency, channel mix, pricing, competition, and macro conditions. It drives payback, capital efficiency, and growth quality. Interpret CAC together with retention and margin, not on its own. The common pitfall is celebrating growth while ignoring whether each customer is becoming more expensive to acquire.'),
('9. CLV (Customer Lifetime Value)',
 'CLV estimates the gross-profit value expected from a customer relationship over time. In the model, CLV is produced through a Monte Carlo framework that reflects uncertainty in pricing, retention, expansion, contraction, and margin.',
 'CLV is driven by ARPA, gross margin, churn behavior, expansion, and discounting. It drives the CLV/CAC relationship and helps users think about long-term customer quality. The pitfall is to treat CLV as exact. It is much more useful as a range that shows whether unit economics are structurally healthy or weak.'),
('10. Net New Subscriptions',
 'Net new subscriptions or net new customers represent the growth in the active base after accounting for both additions and churn. In other words, it is not enough to look at gross additions; the model also shows what was lost and what returned.',
 'Net new is driven by acquisition, churn, reactivation, and timing lags. It drives future ARR and the composition of the customer base. A common forecasting pitfall is reporting strong new-logo activity while hiding that churn erased most of the gain. Net new reveals the true change in the active portfolio.'),
('11. Net Churn',
 'Net churn reflects revenue or customer loss after offsets such as reactivation and expansion are considered. It is a more decision-useful metric than gross churn because it reflects what the business actually lost after recovery effects.',
 'Net churn is driven by gross churn, reactivation, expansion, contraction, and pricing outcomes. It drives confidence in revenue durability and helps explain why the same acquisition level can produce very different growth outcomes. A forecasting pitfall is to quote gross churn alone and miss the recovery side of the system.'),
('12. Expansion',
 'Expansion represents additional revenue generated from the retained base through upsell, cross-sell, seat growth, price realization, or plan upgrades. In the workbook it is separated from contraction so both good and bad monetization dynamics remain visible.',
 'Expansion is driven by account health, product adoption, packaging, pricing, customer maturity, and success motions. It drives NRR and the rate at which the installed base compounds. A common pitfall is to treat expansion as a plug that makes the forecast work. In a good model it must be tied to the retained base and real operating logic.'),
('13. Retention',
 'Retention is the percentage of customers or revenue that remains active over time. In the workbook, retention sits at the center of the recurring-revenue system because it determines whether acquired growth has lasting value.',
 'Retention is driven by product value, onboarding quality, customer success, competition, contract design, and price experience. It drives renewals, expansion opportunity, CAC efficiency, and overall forecast stability. The pitfall is to use one static retention assumption for all cohorts. The model separates customer states and timing to avoid that simplification.'),
('14. Seats / Licenses',
 'Seats or licenses provide a practical scaling unit for SaaS accounts. Even when contracts are priced at the account level, seats or licenses help explain ACV, expansion, and customer mix because they reflect how much product is deployed inside an account.',
 'Seats or licenses are driven by customer size, package structure, product adoption, and expansion success. They drive ACV, TCV, ARR growth, and how pricing changes are felt by the customer. The pitfall is to treat seat count as pure growth. If seat expansion happens at heavy discounting or poor retention, revenue quality may still weaken.'),
]
for title, p1, p2 in metrics:
    doc.add_heading(title, level=2)
    doc.add_paragraph(p1)
    doc.add_paragraph(p2)
    # short user-focused bullet summary
    bullets = []
    if 'ACV' in title: bullets = ['Definition: Annualized contract value per account.', 'Watch: pricing changes, seat mix, segmentation.', 'Pitfall: assuming higher ACV does not affect win rate or churn.']
    elif 'TCV' in title: bullets = ['Definition: total contract value over assumed term.', 'Watch: contract duration and enterprise mix.', 'Pitfall: using TCV instead of recurring value to judge health.']
    elif title.startswith('3. ARR'): bullets = ['Definition: annual recurring revenue.', 'Watch: renewals, new-logo ARR, expansion, contraction, reactivation.', 'Pitfall: forecasting ARR directly without driver logic.']
    elif title.startswith('4. MRR'): bullets = ['Definition: ARR divided by twelve.', 'Watch: monthly trend changes and timing lags.', 'Pitfall: using monthly noise without understanding underlying drivers.']
    elif 'Subscription Bookings' in title: bullets = ['Definition: recurring value booked in the period.', 'Watch: pipeline, close rate, sales capacity, deal timing.', 'Pitfall: confusing booked value with immediately recognized revenue.']
    elif 'Renewals' in title: bullets = ['Definition: retained customer base after churn.', 'Watch: cohort quality, renewal timing, competitive pressure.', 'Pitfall: focusing on renewal counts while ignoring retained value.']
    elif 'Reactivations' in title: bullets = ['Definition: churned customers that return.', 'Watch: timing curves and win-back performance.', 'Pitfall: treating reactivations as brand-new acquisition.']
    elif 'CAC' in title: bullets = ['Definition: customer acquisition cost.', 'Watch: channel efficiency, competition, pricing.', 'Pitfall: celebrating growth while CAC deteriorates.']
    elif 'CLV' in title: bullets = ['Definition: gross-profit lifetime value.', 'Watch: retention, margin, expansion, discounting.', 'Pitfall: treating CLV as precise instead of probabilistic.']
    elif 'Net New' in title: bullets = ['Definition: additions minus losses plus recoveries.', 'Watch: both gross adds and churn.', 'Pitfall: reporting gross growth when net growth is weak.']
    elif 'Net Churn' in title: bullets = ['Definition: loss after offsets such as reactivation and expansion.', 'Watch: gross churn, recovery, expansion, contraction.', 'Pitfall: relying only on gross churn.']
    elif 'Expansion' in title: bullets = ['Definition: added revenue from the retained base.', 'Watch: upsell, pricing, seat growth, adoption.', 'Pitfall: using expansion as a plug.']
    elif 'Retention' in title: bullets = ['Definition: share of customers or revenue that stays.', 'Watch: onboarding, success, value delivery, pricing.', 'Pitfall: using a single static rate for all cohorts.']
    else: bullets = ['Definition: scaling unit within customer accounts.', 'Watch: package mix and account growth.', 'Pitfall: equating seat growth with profitable growth.']
    add_bullets(doc, bullets)

# 4 SaaS user guide
doc.add_heading('4. SaaS Model User Guide', level=1)
doc.add_heading('4.1 Model Overview', level=2)
doc.add_paragraph(
    'The SaaS model forecasts the evolution of a recurring-revenue business over a 24-month horizon. It tracks customer counts, new-logo activity, renewals, churn, reactivation, expansion, contraction, ARR, MRR, bookings, gross margin, CAC payback, and other unit-economics outputs. Its purpose is to help users understand not only where revenue may land, but why.'
)
doc.add_paragraph(
    'The model is especially useful when you want to test whether growth is realistic. For example, can your current sales capacity and implementation capacity support the number of new customers required by your plan? If you increase prices, will that raise ARR enough to offset lower conversion or slightly higher churn? If enterprise deals slip by a quarter, how much revenue moves out of the year? The workbook is built to answer those practical planning questions.'
)
doc.add_paragraph(
    'The key assumptions include starting ARR, starting ARPA, acquisition rate, sales-cycle lag, onboarding lag, capacity limits, churn rates, reactivation behavior, expansion and contraction rates, pricing growth, gross margin, and CAC ratio. Most users should start with the default structure, replace benchmark values with company-specific data where possible, and then build scenarios around the most uncertain assumptions.'
)

doc.add_heading('4.2 Getting Started', level=2)
doc.add_paragraph('Open the workbook and begin on the Dashboard sheet. The dashboard provides the quickest summary of headline outputs, scenario context, and practical notes. Then move to the Assumptions sheet and Parameter Registry to understand which inputs control the model. From there, use the SaaS Model sheet for monthly detail, Monthly Forecast for summary trend review, Sensitivity Analysis for driver ranking, and the Manual Override Log if you make deliberate changes outside the standard input process.')
doc.add_picture(str(sheet_map_png), width=Inches(6.8))
doc.add_paragraph('Figure 4. Suggested workbook navigation order.')
add_bullets(doc, [
    'Start on Dashboard to understand headline outputs and warnings.',
    'Review Assumptions and Parameter Registry before changing any numbers.',
    'Read the Calculations sheet if you need to understand bridge logic.',
    'Use the SaaS Model sheet for monthly customer and revenue mechanics.',
    'Use Sensitivity Analysis to identify which assumptions matter most.'
])

doc.add_heading('4.3 Setting Your Assumptions', level=2)
doc.add_paragraph('The best way to customize the SaaS model is to update the Parameter Registry with your own starting values and assumption ranges. Replace starting ARR and ARPA with your current recurring-revenue profile. Update acquisition rate using recent new-logo history. Set sales-cycle and onboarding lags to reflect your actual timeline from pipeline creation to recognized recurring revenue. Adjust capacity assumptions to reflect how many deals your sales and implementation teams can realistically process in a month.')
doc.add_paragraph('If you know your renewal behavior by segment, adjust retained churn, new-customer churn sensitivity, and reactivated-customer churn sensitivity so they resemble your business. If you have reliable pricing data, update annual price growth and elasticity assumptions. If you know your actual CAC and gross margin, replace benchmark-based settings with observed values. The more company-specific your inputs, the more decision-useful the forecast becomes.')
doc.add_paragraph('Recommended starting point: use your latest quarter or last twelve months as the anchor for starting ARR, ARPA, acquisition rate, gross margin, and realized churn. Then use the model to test uncertainty rather than to recreate perfect history. The workbook is strongest when it is used to explore the next set of management decisions, not to overfit past noise.')
add_bullets(doc, [
    'Use one trusted source for each starting assumption whenever possible.',
    'Change only a few assumptions at a time so cause and effect stay clear.',
    'Document every non-standard edit in the Manual Override Log.',
    'Avoid mixing aspirational targets with observed baseline data.'
])

doc.add_heading('4.4 Running Scenarios', level=2)
doc.add_paragraph('The model is built for scenario planning, not just for a single base case. Start with a base case that reflects your most likely operating path. Then run a bull case to test upside from better acquisition, stronger retention, cleaner renewals, or better pricing realization. Next, run downside cases such as recession, competition, or saturation to see how sensitive the business is to slower growth or weaker economics. Finally, run stress cases like failed pricing, deal slippage, or top-customer loss when leadership needs to understand downside exposure.')
doc.add_paragraph('When interpreting scenarios, focus less on the final ARR number by itself and more on the path the system took to get there. Did ARR fall because acquisition slowed? Did payback worsen because CAC rose? Did retention hold but expansion weaken? Did implementation capacity stop the business from converting demand into revenue quickly enough? The path usually tells you what action management should take.')
add_df_table(doc, saas_scen[['Scenario','Month24 ARR','Average Retention','Average NRR','Average CAC Payback']], 'Illustrative SaaS scenario outputs')

doc.add_heading('4.5 Understanding the Outputs', level=2)
doc.add_paragraph('The monthly outputs show the full customer and revenue bridge month by month. Look at Beginning Customers, Recognized New Logo Customers, Gross Logo Churn Customers, Reactivated Customers, Ending Customers, Beginning ARR, New Logo ARR, Churned ARR, Gross Expansion ARR, Contraction ARR, and Ending ARR. Together they show exactly how recurring revenue is changing.')
doc.add_paragraph('The ARR bridge is particularly useful for management discussions because it separates the four most important questions: how much recurring revenue you kept, how much you lost, how much new recurring revenue you created, and how much additional value you expanded inside the installed base. If growth is disappointing, the bridge helps you identify whether the issue is acquisition, retention, monetization, timing, or capacity.')
doc.add_paragraph('The unit-economics outputs add a second layer of interpretation. CAC payback tells you how long it takes to recover acquisition spend through gross profit. The Rule of 40 proxy provides a quick test of the trade-off between growth and efficiency. CLV/CAC is presented as a range, not as a single truth, because uncertainty matters. Use these outputs together: strong top-line growth with worsening payback is a warning sign, not necessarily a success story.')
add_df_table(doc, mc, 'Illustrative SaaS unit-economics uncertainty summary')
doc.add_picture(str(OUT / 'saas_v4_scenarios.png'), width=Inches(6.8))
doc.add_paragraph('Figure 5. SaaS scenario range over 24 months.')
doc.add_picture(str(OUT / 'saas_v4_driver_tornado.png'), width=Inches(6.8))
doc.add_paragraph('Figure 6. Example of dominant-driver ranking for SaaS planning.')

doc.add_heading('4.6 Practical Applications', level=2)
for title, text in [
    ('Annual planning', 'Use the base case to build the operating plan, then use the downside scenarios to test whether the plan is still viable if acquisition slows or renewals soften. This is especially helpful for setting realistic growth expectations and contingency plans.'),
    ('Fundraising', 'Investors often want to see that management understands both upside and downside. The model helps you show a credible range, explain the mechanics behind the range, and discuss what levers management can control.'),
    ('Board reporting', 'The monthly bridges and driver rankings make board conversations more productive because the focus shifts from “why did the forecast change?” to “which drivers changed and what is management doing about them?”'),
    ('Resource planning', 'Sales capacity and implementation throughput can be translated into hiring implications. If the forecast requires more recognized new-logo activity than the current capacity allows, the model makes that bottleneck visible.'),
    ('Pricing decisions', 'Before changing price, test how much additional ACV or ARPA you gain, how much churn risk rises, and whether the net effect actually strengthens ARR and unit economics.')
]:
    doc.add_heading(title, level=3); doc.add_paragraph(text)

doc.add_heading('4.7 Advanced Features', level=2)
add_bullets(doc, [
    'Tornado analysis: use the dominant-driver output to rank which assumptions deserve management attention first.',
    'Capacity constraints: the model can show when demand exists but cannot be converted into recognized revenue because sales or implementation throughput is limited.',
    'Enterprise lumpiness: large deals do not arrive smoothly every month; the model allows lumpy timing so monthly results look more like real operating conditions.',
    'Cohort aging: newer customers can behave differently from retained or reactivated customers, which makes retention behavior more realistic.',
    'Feedback loops: slower growth can reduce future expansion opportunity, while weaker retention increases the acquisition burden required to support the same ARR outcome.'
])
doc.add_paragraph('A useful advanced workflow is to combine the dominant-driver view with scenario analysis. First identify the top two or three variables that move ARR the most. Then build scenarios around those variables. That approach usually produces better management conversations than changing ten assumptions at once.')

# 5 Streaming guide
doc.add_heading('5. Streaming Model User Guide', level=1)
doc.add_heading('5.1 Model Overview', level=2)
doc.add_paragraph('The Streaming model forecasts subscriber and revenue dynamics over a 24-month horizon. It tracks beginning subscribers, new adds, gross churn, voluntary and involuntary churn, reactivations, ending subscribers, weighted ARPU, billing leakage, effective MRR, ARR run rate, and the effect of price and competitive pressure. The purpose is to show how recurring media revenue changes as acquisition, churn, reactivation, and monetization interact.')
doc.add_paragraph('This model is most useful for subscription media, content businesses, and membership platforms where customer counts and monetization quality both matter. It can help teams think about subscriber growth, reactivation strategy, price increases, leakage reduction, competitive response, and the revenue effect of different acquisition environments. Like the SaaS model, it is most valuable when used as a planning and scenario tool rather than as a single-point forecast.')
doc.add_paragraph('Important assumptions include starting rate-card MRR, subscriber-weighted ARPU, acquisition rate, market-capacity guardrails, price growth, price elasticity, billing leakage, reactivation curves, expansion and contraction rates, and scenario conditions such as competition or recession pressure.')

doc.add_heading('5.2 Getting Started', level=2)
doc.add_paragraph('Open the workbook and review the Dashboard first. Then move to the Assumptions sheet and Parameter Registry. The Streaming Model sheet provides the monthly subscriber and revenue bridge. The Reactivation Curves sheet shows how win-back effects occur over time. Sensitivity Analysis highlights which drivers matter most, and the Dependency Map reminds users which factors tend to move together.')
add_bullets(doc, [
    'Check the starting subscriber base and weighted ARPU first.',
    'Review billing leakage before interpreting effective revenue.',
    'Use the Streaming Model sheet for the detailed monthly bridge.',
    'Use Sensitivity Analysis to identify whether acquisition, pricing, or churn is dominating the outlook.'
])

doc.add_heading('5.3 Setting Your Assumptions', level=2)
doc.add_paragraph('To customize the Streaming model, begin by replacing the starting rate-card MRR and weighted ARPU with your platform data. If your service has multiple plans, use a weighted starting ARPU that reflects actual subscriber mix rather than a simple average list price. Update acquisition rate based on recent gross-add performance, and update billing leakage using observed payment-failure and recovery data if you have it.')
doc.add_paragraph('Next, review your churn and reactivation assumptions. If your platform has separate patterns for new, retained, and reactivated subscribers, preserve that distinction. If you have evidence that win-back campaigns are stronger or weaker than the default assumption, adjust the reactivation curve thoughtfully. Then review price growth and elasticity assumptions, especially if you are planning a tier change, ad-supported expansion, or promotional strategy shift.')
doc.add_paragraph('A good starting point is to use your last three to six months of operating data to anchor ARPU, churn, new adds, and leakage. Then test scenarios around the assumptions that are least certain: pricing response, acquisition efficiency, churn pressure, competition, and recovery effectiveness.')

doc.add_heading('5.4 Running Scenarios', level=2)
doc.add_paragraph('Start with a base case reflecting your most likely path. Then run a bull case to test stronger acquisition, better price realization, or cleaner leakage control. Use recession, competition, and saturation scenarios to test downside subscriber and monetization pressure. Use a failed-price scenario when you want to understand whether a price increase might backfire. Finally, use a severe stress scenario to understand the lower boundary of the model under combined pressure.')
doc.add_paragraph('When reading scenarios, avoid focusing only on ending ARR. Look at how the monthly path changes. A scenario can finish only modestly below the base case while still creating a much rougher operating path, sharper net churn, or weaker reactivation. Those pattern changes often matter more for decision-making than the final point estimate.')
add_df_table(doc, stream_scen[['Scenario','Month24 ARR','Average Net Churn','Average Dampening','Min ARR']], 'Illustrative Streaming scenario outputs')

doc.add_heading('5.5 Understanding the Outputs', level=2)
doc.add_paragraph('The Streaming model provides a monthly subscriber and revenue bridge. Begin with Beginning Subscribers, Recognized New Adds, Gross Churn, Total Reactivations, and Ending Subscribers. Then move to the revenue side: Rate Card MRR, Effective MRR after leakage, Gross Expansion MRR, Contraction MRR, Total MRR, and ARR Run Rate. This flow shows why subscriber growth and revenue growth do not always move in lockstep.')
doc.add_paragraph('Churn analysis is especially important. Gross churn tells you how many subscribers were lost before recovery. Net churn shows the loss after reactivations offset part of the damage. Reactivation Dampening tells you how much of gross churn is being absorbed by returns. A platform with mediocre gross churn may still stabilize well if reactivation is strong. A platform with low reported churn may still be weaker than it appears if leakage or contraction is high.')
doc.add_paragraph('ARPU dynamics deserve separate attention. The workbook allows pricing and monetization changes to affect revenue through weighted ARPU, while leakage and contraction prevent the model from overstating collected revenue. This is useful when evaluating tier migration, promotional strategies, or subscription-package changes. Revenue quality depends not only on list price but also on what is actually collected and retained.')
doc.add_picture(str(OUT / 'streaming_v4_scenarios.png'), width=Inches(6.8))
doc.add_paragraph('Figure 7. Streaming scenario range over 24 months.')
doc.add_picture(str(OUT / 'streaming_v4_driver_tornado.png'), width=Inches(6.8))
doc.add_paragraph('Figure 8. Example of dominant-driver ranking for Streaming planning.')

doc.add_heading('5.6 Practical Applications', level=2)
for title, text in [
    ('Content investment', 'Use the model to ask whether a higher acquisition push actually leads to durable subscriber and revenue growth, or whether weak retention causes the spend to fade quickly.'),
    ('Pricing strategy', 'Test how a price increase affects ARPU, acquisition, churn, and collected revenue. A pricing change that looks attractive on a static spreadsheet may be unattractive once elasticity and leakage are included.'),
    ('Churn reduction', 'Use churn and reactivation outputs to identify whether retention interventions should focus on new subscribers, retained subscribers, or recovery of recently churned users.'),
    ('Reactivation campaigns', 'Estimate whether a win-back campaign is worth funding by looking at how stronger reactivation changes net churn and downstream MRR.'),
    ('Subscriber acquisition planning', 'The model can show whether you need better top-of-funnel efficiency, better retention, or both. It is particularly useful when management asks whether higher acquisition can solve a retention problem. Often it cannot.'),
]:
    doc.add_heading(title, level=3); doc.add_paragraph(text)

doc.add_heading('5.7 Advanced Features', level=2)
add_bullets(doc, [
    'Billing leakage: use this to test how payment failure and recovery affect collected recurring revenue.',
    'Cohort segmentation: newer, retained, and reactivated subscribers can behave differently over time.',
    'Promotional timing: use scenario assumptions to test whether near-term boosts create later churn or revenue dilution.',
    'Tier migration: interpret ARPU changes as a mix of price, product tier, and subscriber behavior rather than a static list-price change.',
    'Competitive response: use competition and stress scenarios to estimate what happens if switching pressure increases or price realization weakens.'
])
doc.add_paragraph('For advanced users, one of the best habits is to compare gross churn, net churn, reactivation dampening, and effective MRR in the same review. That combination often reveals whether the business is really solving subscriber loss or merely covering it with short-term acquisition or pricing actions.')

# 6 Best practices
doc.add_heading('6. Best Practices and Tips', level=1)
add_bullets(doc, [
    'Start with the base case, then move to scenarios. A clean base case provides the anchor that makes scenario interpretation useful.',
    'Update assumptions quarterly at minimum, and monthly for the most volatile drivers such as acquisition, churn, ARPU, pricing response, or CAC.',
    'Document all material changes in the Manual Override Log so future users understand why the forecast changed.',
    'Compare forecast outputs to actuals. The biggest forecasting gains usually come from learning which assumptions repeatedly drift away from reality.',
    'Use dominant-driver analysis to narrow management attention to the variables that truly move the result.',
    'Do not over-optimize precision. The workbook is strongest when it shows direction, magnitude, and trade-offs, not false certainty.',
    'Communicate uncertainty explicitly to stakeholders. Present the base case together with downside and upside cases.',
    'Combine the model with qualitative judgment. The workbook helps structure thinking, but management should still apply market context, product knowledge, and commercial judgment.'
])
doc.add_paragraph('A particularly useful practice is to separate “known updates” from “judgment updates.” Known updates are items supported by recent actuals, such as last month\'s churn or the current pipeline capacity. Judgment updates are items like the likely effect of a proposed price change or a possible competitive launch. Recording that distinction helps teams understand which parts of the forecast are measurement and which parts are management opinion.')
doc.add_paragraph('Another strong habit is to run one scenario specifically for stakeholder communication. For example, for a board meeting you might present base, downside, and stress. For internal budget setting you might present base, hiring-constrained, and pricing-risk scenarios. For fundraising you may want base, upside, and downside. The best scenario set depends on the decision you are trying to support.')

# 7 FAQ
doc.add_heading('7. Common Questions and Troubleshooting', level=1)
faqs = [
('What if my company does not match the archetype?', 'Replace the starting assumptions with your own operating data as early as possible. The structure of the model is portable even when the default benchmark values are not.'),
('How do I customize the model for my business?', 'Start by updating starting revenue, ARPU or ARPA, acquisition rate, churn, pricing assumptions, and capacity limits. Then adjust scenario ranges around the variables that are most uncertain for your business.'),
('What if I do not have all the data?', 'Use the default assumptions as placeholders, but mark which ones are estimated. Then prioritize collecting the inputs that matter most: acquisition, churn, pricing, gross margin, and capacity.'),
('How often should I update the model?', 'Monthly for fast-changing drivers, quarterly for most strategic assumptions, and immediately after major pricing, competitive, or staffing changes.'),
('Can I use this for multi-year planning?', 'Yes, but use caution. The 24-month horizon is useful for scenario planning, while longer-range planning should be supported by broader strategic assumptions and more uncertainty.'),
('What are the model\'s limitations?', 'These models simplify reality. They are best used to structure thinking and compare scenarios, not to imply certainty. They also depend heavily on assumption quality.'),
('When should I not use these models?', 'Do not use them as exact commitment tools, legal disclosure documents, or substitutes for management judgment. They are especially weak if your business has no reliable starting data or if the revenue model differs fundamentally from recurring revenue.')
]
for q, a in faqs:
    doc.add_heading(q, level=2)
    doc.add_paragraph(a)

# 8 Appendices
doc.add_heading('8. Appendices', level=1)
doc.add_heading('Appendix A. Glossary of Terms', level=2)
gloss = pd.DataFrame([
    ['ARR','Annual recurring revenue'], ['MRR','Monthly recurring revenue'], ['ACV','Annual contract value'], ['TCV','Total contract value'], ['CAC','Customer acquisition cost'], ['CLV','Customer lifetime value'], ['NRR','Net revenue retention'], ['ARPU','Average revenue per user'], ['ARPA','Average revenue per account'], ['Reactivation','Return of a previously churned customer or subscriber'], ['Billing leakage','Revenue lost because billed value is not fully collected'], ['Dominant drivers','The assumptions that move the forecast the most'],
], columns=['Term','Meaning'])
add_df_table(doc, gloss)

doc.add_heading('Appendix B. Formula Reference Guide', level=2)
formulas = pd.DataFrame([
    ['SaaS','Ending ARR','Retained ARR After Churn + Reactivated ARR + New Logo ARR + Gross Expansion ARR - Contraction ARR - Top Customer Shock ARR'],
    ['SaaS','Ending MRR','Ending ARR / 12'],
    ['SaaS','Retention Rate','Renewals / Beginning Customers'],
    ['SaaS','CAC Payback','CAC per New Customer / (ARPA × Gross Margin)'],
    ['SaaS','Rule of 40 Proxy','Annualized growth rate + adjusted gross margin proxy'],
    ['Streaming','Ending Subscribers','Beginning Subscribers - Gross Churn + Recognized New Adds + Total Reactivations'],
    ['Streaming','Effective MRR','Rate Card MRR × (1 - Billing Leakage %)'],
    ['Streaming','ARR Run Rate','Total MRR × 12'],
    ['Streaming','Net Churn Rate','(Gross Churn - Total Reactivations) / Beginning Subscribers'],
    ['Streaming','Reactivation Dampening %','Total Reactivations / Gross Churn'],
], columns=['Model','Metric','Formula'])
add_df_table(doc, formulas)

doc.add_heading('Appendix C. Data Source List', level=2)
add_df_table(doc, source[['Source_ID','Element','Status','Methodology']])

doc.add_heading('Appendix D. Parameter Registry Template', level=2)
param_template = param[['Domain','Parameter','Value','Unit','Type','Owner','Description']].copy()
add_df_table(doc, param_template, max_rows=20)

doc.add_heading('Appendix E. Scenario Planning Worksheet', level=2)
scenario_ws = pd.DataFrame([
    ['Question','Example prompt'],
    ['What decision are we making?','Should we increase price next quarter?'],
    ['What is the base case?','Current acquisition, churn, and capacity assumptions'],
    ['What is the upside case?','Higher conversion, better expansion, lower churn'],
    ['What is the downside case?','Competitive pressure, higher churn, slower pipeline'],
    ['Which 3 drivers matter most?','Acquisition, retention, pricing response'],
    ['What would management do if downside occurs?','Cut spend, revise hiring, adjust price strategy'],
], columns=['Worksheet Step','How to Use It'])
add_df_table(doc, scenario_ws)

doc.add_paragraph('Final reminder: use these models as working forecasting systems. Update them regularly, challenge them openly, compare them to actuals, and use them to support better decisions—not just to produce a number.')

doc.save(GUIDE)
print(f'Created: {GUIDE}')
